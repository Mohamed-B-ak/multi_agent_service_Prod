from fpdf import FPDF
import openpyxl
from docx import Document
from crewai.tools import BaseTool
class TextToPDFTool(BaseTool):
    name: str = "Text to PDF Tool"
    description: str = "Converts a text string to a PDF file."

    def _run(self, text_content: str, output_path: str) -> str:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        for line in text_content.splitlines():
            pdf.cell(200, 10, txt=line.strip(), ln=True)

        pdf.output(output_path)
        return f"✅ PDF created: {output_path}"


class TextToWordTool(BaseTool):
    name: str = "Text to Word Tool"
    description: str = "Converts a text string to a Word (.docx) file."

    def _run(self, text_content: str, output_path: str) -> str:
        document = Document()
        document.add_heading("Generated from Text", level=1)

        for line in text_content.splitlines():
            document.add_paragraph(line.strip())

        document.save(output_path)
        return f"✅ Word document created: {output_path}"


class TextToExcelTool(BaseTool):
    name: str = "Text to Excel Tool"
    description: str = "Converts a text string to an Excel (.xlsx) file."

    def _run(self, text_content: str, output_path: str) -> str:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "From Text"

        for idx, line in enumerate(text_content.splitlines(), start=1):
            sheet.cell(row=idx, column=1, value=line.strip())

        workbook.save(output_path)
        return f"✅ Excel file created: {output_path}"