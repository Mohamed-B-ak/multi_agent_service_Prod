import os
from crewai import LLM

# 🔐 Set your OpenAI API Key directly (for local testing only)
os.environ["OPENAI_API_KEY"] = "*************************"  # <-- Put your actual API key here

llm_obj = LLM(
    model="gpt-3.5-turbo",
    api_key=os.getenv("OPENAI_API_KEY"),  # This now resolves properly
    temperature=0.3,
)








from fpdf import FPDF
from crewai.tools import BaseTool

class TextToPDFTool(BaseTool):
    name: str = "Text to PDF Tool"
    description: str = "Converts a text string to a PDF file."

    def _run(self, text_content: str, output_path: str) -> str:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        for line in text_content.splitlines():
            pdf.cell(200, 10, txt=line.strip(), ln=True)

        pdf.output(output_path)
        return f"✅ PDF created: {output_path}"
from docx import Document
from crewai.tools import BaseTool

class TextToWordTool(BaseTool):
    name: str = "Text to Word Tool"
    description: str = "Converts a text string to a Word (.docx) file."

    def _run(self, text_content: str, output_path: str) -> str:
        document = Document()
        document.add_heading("Generated from Text", level=1)

        for line in text_content.splitlines():
            document.add_paragraph(line.strip())

        document.save(output_path)
        return f"✅ Word document created: {output_path}"
import openpyxl
from crewai.tools import BaseTool

class TextToExcelTool(BaseTool):
    name: str = "Text to Excel Tool"
    description: str = "Converts a text string to an Excel (.xlsx) file."

    def _run(self, text_content: str, output_path: str) -> str:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "From Text"

        for idx, line in enumerate(text_content.splitlines(), start=1):
            sheet.cell(row=idx, column=1, value=line.strip())

        workbook.save(output_path)
        return f"✅ Excel file created: {output_path}"
    

from crewai import Agent, LLM, Task, Crew

def file_creation_ageent(llm_obj) -> Agent:
    """
    Agent that creates PDF, Word, and Excel files from provided text content.

    Args:
        llm_obj: The LLM instance to use for task understanding.

    Returns:
        Agent instance ready to use in a Crew.
    """
    return Agent(
        role="File Creator",
        goal="Generate user-requested files from text in the desired format",
        backstory=(
            "You are a helpful assistant capable of generating well-structured documents "
            "from raw user input. You support creating PDF, Word (.docx), and Excel (.xlsx) files. "
            "Always confirm the file format and save location."
        ),
        tools=[
            TextToPDFTool(),
            TextToWordTool(),
            TextToExcelTool()
        ],
        allow_delegation=False,
        llm=llm_obj,
        verbose=True
    )





if __name__ == "__main__":
    import os
    # 🔐 Setup LLM
    llm_obj = LLM(
        model="gpt-3.5-turbo",
        api_key=os.getenv("OPENAI_API_KEY"),
        temperature=0.3,

    )

    # 🧠 Create agent
    
    from agents.file_creation_agent import file_creation_agent
    agent = file_creation_agent(llm_obj)
    # 🧾 Sample input
    sample_text = """Name, Age, Country
                     Alice, 30, USA
                     Bob, 25, Canada
                     Charlie, 28, UK"""

    user_prompt = "create a word file "

    # 🧩 Create task
    task = Task(
        description=f"Create a file based on he user prompt {user_prompt} with the following text content :\n{sample_text}",
        expected_output="A confirmation message that the file was created.",
        agent=agent
    )

    # 🚀 Run the Crew
    crew = Crew(agents=[agent], tasks=[task], verbose=True)
    result = crew.kickoff()

    print("\n✅ Final Result:\n", result)